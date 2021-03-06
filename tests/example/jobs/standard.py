import time
import asyncio
import threading
from datetime import timedelta

import greenlet

from pq import api


class TestError(Exception):
    pass


class RunPyCode(api.Job):
    '''execute python code in *code*. There must be a *task_function*
function defined which accept key-valued parameters only.'''
    timeout = timedelta(seconds=60)

    def __call__(self, code=None, **kwargs):
        code_local = compile(code, '<string>', 'exec')
        ns = {}
        exec(code_local, ns)
        func = ns['task_function']
        return func(**kwargs)


class Addition(api.Job):
    timeout = timedelta(seconds=60)

    def __call__(self, a=0, b=0):
        return a + b


class Asynchronous(api.Job):

    async def __call__(self, lag=1):
        start = time.time()
        await asyncio.sleep(lag)
        return time.time() - start


@api.job()
async def notoverlap(self, lag=1):
    async with self.lock():
        start = time.time()
        await asyncio.sleep(lag)
        return {
            'start': start,
            'end': time.time()
        }


@api.job()
async def queue_from_task(self):
    task = await self.queue_task('asynchronous')
    return task.tojson()


class WorkerInfo(api.Job):

    def __call__(self):
        return self.backend.info()


class GreenExecutor(api.Job):

    def __call__(self):
        return self.run_in_executor(self.backend.info)


class CpuBound(api.Job):
    concurrency = api.CPUBOUND

    def __call__(self, error=False):
        self.logger.info('Testing CpuBound concurrency')
        self.logger.warning('Sleeping for 2 seconds')
        time.sleep(1)
        if error:
            raise TestError('just a test')
        return ['OK', 2]


@api.job()
def testlocalqueue(self):
    return self.backend.queues()


class CpuBoundWithAsync(api.Job):
    concurrency = api.CPUBOUND

    def __call__(self, asyncio=False):
        if asyncio:
            return self.asyncio()
        else:
            return self.greenlet_info()

    def greenlet_info(self):
        return greenlet.getcurrent().parent is not None

    async def asyncio(self):
        await asyncio.sleep(1)
        return self.greenlet_info()


class CpuBoundBigLog(api.Job):
    concurrency = api.CPUBOUND

    def __call__(self):
        # Log more date then the pipe buffer, as logs are send through the pipe
        for i in range(1024):
            self.backend.logger.debug('*'*1024)


@api.job()
async def scrape(self, url=None):
    assert url, "url is required"
    request = await self.http.get(url)
    return request.text()


@api.job(concurrency=api.THREAD_IO)
def extract_docx(self, input=None, output=None):
    """
    Extract text from a docx document

    This task is not async friendly and therefore it should be run as
    THREAD_IO or as CPUBOUND

    :return: the length of the text extracted
    """
    import docx
    assert input and output, "input and output must be given"
    document = docx.Document(input)
    text = '\n\n'.join(_docx_text(document))
    with open(output, 'w') as fp:
        fp.write(text)
    return {
        'thread': threading.get_ident(),
        'text': len(text)
    }


def _docx_text(document):
    for paragraph in document.paragraphs:
        yield paragraph.text

    yield from _docx_tables(document.tables)


def _docx_tables(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                # For every cell in every row of the table, extract text from
                # child paragraphs.
                for paragraph in cell.paragraphs:
                    yield paragraph.text

                # Then recursively extract text from child tables.
                yield from _docx_tables(cell.tables)
